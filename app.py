# app.py

import pandas as pd
import streamlit as st
import time
import re
import os
import zipfile
import openai
from prompts import prompts
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from html.parser import HTMLParser

def render_expanders(expanders):
    for key, value in expanders.items():
        with st.expander(value["title"]):
            st.markdown(value["content"])

def create_url_path(keyword):
    url_path = keyword.lower()
    url_path = re.sub(r"[^a-z0-9\s]+", "", url_path)  # Remove non-alphanumeric and non-space characters
    url_path = url_path.replace(" ", "-")  # Replace spaces with hyphens

    # Remove trailing hyphen if present
    if url_path[-1] == "-":
        url_path = url_path[:-1]

    return f"/{url_path}"

def create_full_path(domain, url_path):
    return f"https://{domain}{url_path}"

def generate_content(api_key, prompt, sections, model, temperature, presence_penalty, frequency_penalty, max_tokens):
    openai.api_key = api_key

    system_message = prompts["system_message"]
    
    print(f"Generated prompt:\n{prompt}\n")

    messages = [
        {"role": "system", "content": system_message},
        {"role": "user", "content": prompt}
    ]
    
    
    print(f"API Call Parameters:")
    print(f"API Key: {api_key}")
    print(f"GPT Model: {model}")
    print(f"Prompt: {prompt}")
    print(f"Sections: {sections}")
    print(f"Temperature: {temperature}")
    print(f"Presence Penalty: {presence_penalty}")
    print(f"Frequency Penalty: {frequency_penalty}")
    print(f"Max Tokens: {max_tokens}")

    completion = openai.ChatCompletion.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        presence_penalty=presence_penalty,
        frequency_penalty=frequency_penalty,
        messages=messages,
    )

    response = completion.choices[0].message.content.strip()
    # print(f"Generated response:\n{response}\n")
    return response

def generate_related_links(df, current_topic):
    current_category = df.loc[df['topic'] == current_topic, 'category'].values[0]
    current_full_path = df.loc[df['topic'] == current_topic, 'full path'].values[0]
    related_links = df[df['category'] == current_category][['topic', 'full path']]
    
    # Filter out the self-link by comparing the full paths
    related_links = related_links[related_links['full path'] != current_full_path]

    return related_links.to_dict('records')

def generate_article(api_key, topic, sections, related_links, model, temperature, presence_penalty, frequency_penalty, max_tokens):
    if related_links:
        related_links_prompt = prompts["related_links_prompt"].format(
            ", ".join([f"{rl['topic']} ({rl['full path']})" for rl in related_links])
        )
    else:
        related_links_prompt = ""

    prompt = prompts["article_prompt"].format(
        topic, "\n".join(str(sec) for sec in sections), related_links_prompt
    )

    article = generate_content(api_key, prompt, sections, model, temperature, presence_penalty, frequency_penalty, max_tokens)
    return article

def generate_docx(topic, article, sections, file_name):
    doc = Document()

    # Set up the document title
    doc.add_heading(topic, level=1)

    # Add the content to the document
    parser = HTMLParser()
    for section, content in zip(sections, article.split("\n\n")):
        doc.add_heading(section, level=2)
        parser.feed(content)
        for line in parser.text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        parser.reset()

    # Save the document
    doc.save(file_name)


